import pandas as pd
import re
import os
import docx
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """צביעת רקע של תא בטבלה"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def split_ingredients(ing_str):
    """פירוק מחרוזת המרכיבים לרשימה מופרדת לפי כמויות ומזיגות"""
    pattern = r'(\d+(?:\.\d+)?\s*ml|\d+\s*Bar Spoon|A\s+splash\s+of|Few\s+Drops\s+of|Few\s+drops\s+of|Top\s+with|\d+\s+Dash(?:es)?|\d+\s+Drop(?:s)?)'
    matches = list(re.finditer(pattern, ing_str, re.IGNORECASE))
    
    ingredients = []
    for i in range(len(matches)):
        start = matches[i].start()
        end = matches[i+1].start() if i + 1 < len(matches) else len(ing_str)
        ingredients.append(ing_str[start:end].strip())
    
    if not ingredients:
        return [ing_str]
    return ingredients

def main():
    input_file = 'iba_live_scraped_data.csv'
    output_file = 'iba_cocktails_ingredients.docx'
    
    if not os.path.exists(input_file):
        print(f"Error: Could not find '{input_file}' in your directory.")
        return

    print("Loading dataset and generating Word document...")
    df = pd.read_csv(input_file)
    
    doc = docx.Document()

    # הגדרת העמוד לתצורת רוחב (Landscape) בשביל מקום לעמודות
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    # צמצום שוליים למקסימום מקום
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # כותרת המסמך
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("IBA Cocktails - Ingredients Directory")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = docx.shared.RGBColor(200, 122, 50) # צבע אמבר בר אסתטי

    # תת כותרת
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run("Each ingredient is separated into its own column for clear structured viewing.")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True

    # יצירת טבלה של 9 עמודות (שם + עד 8 מרכיבים)
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'

    # הגדרת כותרות הטבלה
    headers = ["Cocktail Name"] + [f"Ingredient {i}" for i in range(1, 9)]
    hdr_cells = table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "C87A32") # רקע כותרת מוזהב/אמבר
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.color.rgb = docx.shared.RGBColor(255, 255, 255)

    # מעבר על השורות ופיצול לעמודות
    for idx, row in df.iterrows():
        name = row['Cocktail Name']
        ing_str = row['Ingredients']
        if pd.isna(ing_str):
            continue
            
        ings = split_ingredients(ing_str)
        row_cells = table.add_row().cells
        
        # עמודת שם הקוקטייל
        row_cells[0].text = str(name)
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_background(row_cells[0], "FDFBF9") # טינט קרם עדין לשם הקוקטייל
        
        # עמודות המרכיבים המפוצלים
        for i in range(8):
            if i < len(ings):
                row_cells[i+1].text = ings[i]
            else:
                row_cells[i+1].text = ""
            if row_cells[i+1].paragraphs[0].runs:
                row_cells[i+1].paragraphs[0].runs[0].font.name = 'Arial'
                row_cells[i+1].paragraphs[0].runs[0].font.size = Pt(10)

    doc.save(output_file)
    print(f"Success! Created Word file at: {os.path.abspath(output_file)}")

if __name__ == "__main__":
    main()