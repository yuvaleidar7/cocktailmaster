import pandas as pd
import docx
import io

def process_word_to_dataframe(docx_filename):
    doc = docx.Document(docx_filename)
    lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
            
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ').replace('"', '""')
                row_data.append(f'"{cell_text}"')
            lines.append(",".join(row_data))

    start_index = -1
    for i, line in enumerate(lines):
        # מחפש שורה עם לפחות 5 פסיקים
        if line.count(',') >= 5:
            start_index = i
            break
            
    if start_index == -1:
        raise ValueError("שגיאה: לא הצלחתי למצוא מבנה של CSV בתוך קובץ הוורד.")
        
    relevant_lines = lines[start_index:]
    csv_content = "\n".join(relevant_lines)
    
    df = pd.read_csv(io.StringIO(csv_content), encoding='utf-8')
    return df

def merge_cocktail_datasets():
    docx_file = 'iba_cocktails_ingredients.docx'
    csv_file = 'iba_live_scraped_data.csv'
    output_csv = 'cocktalismaster_iba_dataset.csv'
    output_excel = 'cocktalismaster_iba_dataset.xlsx'
    
    try:
        print(f"קורא נתונים מקובץ ה-Word: {docx_file}...")
        df_hebrew = process_word_to_dataframe(docx_file)
        
        print(f"קורא נתונים מקובץ ה-CSV: {csv_file}...")
        df_english = pd.read_csv(csv_file, encoding='utf-8')
        
        print("מאחד את הנתונים אופקית...")
        merged_df = pd.concat([df_english, df_hebrew], axis=1)
        
        # הדפסת ביקורת לטרמינל כדי שתוכל לראות את העמודות!
        print("\n--- רשימת העמודות בקובץ המאוחד ---")
        for col in merged_df.columns:
            print(f"- {col}")
        print("----------------------------------\n")
        
        # שמירה כ-CSV לשימוש עתידי בקוד
        merged_df.to_csv(output_csv, index=False, encoding='utf-8-sig')
        
        # שמירה כ-Excel כדי שתוכל לפתוח ולראות את זה נקי וברור
        merged_df.to_excel(output_excel, index=False)
        
        print(f"הקבצים אוחדו בהצלחה!")
        print(f"נשמר קובץ קוד: {output_csv}")
        print(f"נשמר קובץ לצפייה נוחה: {output_excel}")
        
    except FileNotFoundError as e:
        print(f"שגיאה: לא ניתן למצוא את הקובץ - {e.filename}. ודא שהוא נמצא באותה התיקייה.")
    except Exception as e:
        print(f"התרחשה שגיאה: {e}")

if __name__ == "__main__":
    merge_cocktail_datasets()